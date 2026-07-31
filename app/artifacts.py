from datetime import datetime
from pathlib import Path

from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill

from .workspace import project_root


class ArtifactError(RuntimeError):
    pass


def _report_dir(project):
    try:
        directory = project_root(project.workspace_path, project.attached_file_path) / "agilesprint_reports"
    except Exception as error:
        raise ArtifactError("请先接入一个存在的本地项目文件夹，才能导出项目资料") from error
    directory.mkdir(exist_ok=True)
    return directory


def export_markdown_summary(project, sprint):
    destination = _report_dir(project) / f"{sprint.name}_协作复盘.md"
    lines = [
        f"# {project.name} · {sprint.name} 协作复盘",
        "",
        f"- 迭代目标：{sprint.goal}",
        f"- 当前状态：{sprint.status}",
        f"- 修复轮次：{sprint.repair_round}",
        "",
        "## 任务清单",
    ]
    for task in sprint.tasks:
        lines.append(f"- [{task.status}] {task.title}｜负责人：{task.assignee or '待分配'}｜预估：{task.estimate or '未估算'}")
    lines.extend(["", "## 协作记录"])
    for message in sprint.messages:
        lines.extend([f"### {message.sender}：{message.summary}", message.content, f"- 决策：{message.decision or '无'}", ""])
    destination.write_text("\n".join(lines), encoding="utf-8")
    return destination


def export_word_test_report(project, sprint, test_runs):
    destination = _report_dir(project) / f"{sprint.name}_测试报告.docx"
    document = Document()
    document.add_heading(f"{project.name} · {sprint.name} 测试报告", 0)
    document.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}")
    document.add_heading("迭代目标", level=1)
    document.add_paragraph(sprint.goal or "未填写")
    document.add_heading("测试结果", level=1)
    if not test_runs:
        document.add_paragraph("尚未执行真实 Python 测试。")
    for index, run in enumerate(test_runs, 1):
        document.add_heading(f"第 {index} 次测试：{'通过' if run.status == 'passed' else '失败'}", level=2)
        document.add_paragraph(f"命令：{run.command}")
        document.add_paragraph(run.output or "无输出")
    document.add_heading("测试 Agent 结论", level=1)
    messages = [message for message in sprint.messages if message.sender_role == "tester"]
    for message in messages[-5:]:
        document.add_paragraph(f"{message.summary}：{message.content}")
    document.save(destination)
    return destination


def export_excel_tasks(project, sprint):
    destination = _report_dir(project) / f"{sprint.name}_任务表.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "任务看板"
    headers = ["任务名称", "说明", "负责人", "状态", "优先级", "预估", "退回原因"]
    sheet.append(headers)
    for cell in sheet[1]:
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="2563EB")
    for task in sprint.tasks:
        sheet.append([task.title, task.description, task.assignee, task.status, task.priority, task.estimate, task.return_reason])
    sheet.freeze_panes = "A2"
    for column, width in {"A": 28, "B": 50, "C": 16, "D": 16, "E": 14, "F": 14, "G": 32}.items():
        sheet.column_dimensions[column].width = width
    workbook.save(destination)
    return destination
